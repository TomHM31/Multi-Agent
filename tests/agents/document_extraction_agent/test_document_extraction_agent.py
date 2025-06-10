import unittest
import os
import sys
import json
import pandas as pd
from unittest.mock import patch, mock_open, MagicMock
from PIL import Image
from docx import Document # Import Document for mocking

# Add the project root to sys.path to allow importing modules from the 'agents' package
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..')))

from agents.document_extraction_agent.document_extraction_agent import DocumentExtractionAgent, PDF_TEXT_THRESHOLD

class TestDocumentExtractionAgent(unittest.TestCase):
    """
    Unit tests for the DocumentExtractionAgent.
    """

    def setUp(self):
        """
        Set up test environment before each test.
        """
        self.agent = DocumentExtractionAgent(
            preprocessed_output_path="test_preprocessed.jsonl",
            dead_letter_queue_path="test_dead_letter.log",
            unsupported_queue_path="test_unsupported_files.log"
        )
        # Clean up any previous test output files
        self._cleanup_files()

    def tearDown(self):
        """
        Clean up test environment after each test.
        """
        self._cleanup_files()

    def _cleanup_files(self):
        """Helper to remove test output files."""
        if os.path.exists(self.agent.preprocessed_output_path):
            os.remove(self.agent.preprocessed_output_path)
        if os.path.exists(self.agent.dead_letter_queue_path):
            os.remove(self.agent.dead_letter_queue_path)
        if os.path.exists(self.agent.unsupported_queue_path):
            os.remove(self.agent.unsupported_queue_path)

    def test_normalize_text_whitespace(self):
        """
        Test text normalization for collapsing whitespace.
        """
        text = "  Hello   World! \n\n This is a test.  "
        expected = "Hello World! This is a test."
        self.assertEqual(self.agent._normalize_text(text), expected)

    def test_normalize_text_encoding_issues(self):
        """
        Test text normalization for common encoding issues.
        """
        text = "This is a test with â€™ and â€œquotesâ€ ."
        expected = "This is a test with ' and \"quotes\"."
        self.assertEqual(self.agent._normalize_text(text), expected)

    def test_normalize_text_boilerplate(self):
        """
        Test text normalization for removing boilerplate.
        """
        text = "Document content.\nPage 1 of 10\nConfidential Document"
        expected = "Document content."
        self.assertEqual(self.agent._normalize_text(text), expected)

    def test_normalize_text_empty(self):
        """
        Test text normalization with empty input.
        """
        self.assertEqual(self.agent._normalize_text(""), "")
        self.assertEqual(self.agent._normalize_text(None), "")

    @patch('pandas.read_csv')
    def test_extract_csv_text_single_row(self, mock_read_csv):
        """
        Test CSV extraction for a single row.
        """
        mock_df = pd.DataFrame([
            {"col1": "value1", "col2": "value2", "col3": "value3"}
        ])
        mock_read_csv.return_value = mock_df
        expected_text = "col1: value1\ncol2: value2\ncol3: value3"
        self.assertEqual(self.agent._extract_csv_text("dummy.csv"), expected_text)

    @patch('pandas.read_csv')
    def test_extract_csv_text_multiple_rows(self, mock_read_csv):
        """
        Test CSV extraction for multiple rows.
        """
        mock_df = pd.DataFrame([
            {"colA": "valA1", "colB": "valB1"},
            {"colA": "valA2", "colB": "valB2"}
        ])
        mock_read_csv.return_value = mock_df
        expected_text = "colA: valA1\ncolB: valB1\n\ncolA: valA2\ncolB: valB2"
        self.assertEqual(self.agent._extract_csv_text("dummy.csv"), expected_text)

    @patch('pandas.read_csv')
    def test_extract_csv_text_with_empty_cells(self, mock_read_csv):
        """
        Test CSV extraction with empty cells.
        """
        mock_df = pd.DataFrame([
            {"col1": "value1", "col2": None, "col3": "value3"},
            {"col1": "value4", "col2": "", "col3": "value6"}
        ])
        mock_read_csv.return_value = mock_df
        expected_text = "col1: value1\ncol3: value3\n\ncol1: value4\ncol3: value6"
        self.assertEqual(self.agent._extract_csv_text("dummy.csv"), expected_text)

    @patch('pandas.read_csv', side_effect=Exception("CSV read error"))
    def test_extract_csv_text_failure(self, mock_read_csv):
        """
        Test CSV extraction failure.
        """
        result = self.agent._extract_csv_text("bad.csv")
        self.assertEqual(result, "")
        with open(self.agent.dead_letter_queue_path, "r") as f:
            log_content = f.read()
        expected_log = f"source_id: bad.csv, step: CSV extraction, error: CSV read error\n"
        self.assertIn(expected_log, log_content)

    @patch('pdfplumber.open')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._ocr_pdf')
    def test_extract_pdf_text_success(self, mock_ocr_pdf, mock_pdfplumber_open):
        """
        Test PDF text extraction with sufficient text.
        """
        mock_pdf = MagicMock()
        mock_page1 = MagicMock()
        mock_page1.extract_text.return_value = "This is page one content."
        mock_page2 = MagicMock()
        mock_page2.extract_text.return_value = "This is page two content."
        mock_pdf.pages = [mock_page1, mock_page2]
        mock_pdfplumber_open.return_value.__enter__.return_value = mock_pdf

        text, page_count = self.agent._extract_pdf_text("dummy.pdf")
        self.assertIn("This is page one content.", text)
        self.assertIn("This is page two content.", text)
        self.assertEqual(page_count, 2)
        mock_ocr_pdf.assert_not_called() # OCR should not be called

    @patch('pdfplumber.open')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._ocr_pdf', return_value="OCR text from image.")
    def test_extract_pdf_text_fallback_to_ocr_short_text(self, mock_ocr_pdf, mock_pdfplumber_open):
        """
        Test PDF text extraction falling back to OCR due to short text.
        """
        mock_pdf = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "short" # Text shorter than threshold
        mock_pdf.pages = [mock_page]
        mock_pdfplumber_open.return_value.__enter__.return_value = mock_pdf

        text, page_count = self.agent._extract_pdf_text("dummy.pdf")
        self.assertEqual(text, "OCR text from image.")
        self.assertEqual(page_count, 1)
        mock_ocr_pdf.assert_called_once_with("dummy.pdf")

    @patch('pdfplumber.open', side_effect=Exception("PDF open error"))
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._ocr_pdf', return_value="OCR text from image.")
    def test_extract_pdf_text_fallback_to_ocr_on_error(self, mock_ocr_pdf, mock_pdfplumber_open):
        """
        Test PDF text extraction falling back to OCR on error.
        """
        text, page_count = self.agent._extract_pdf_text("bad.pdf")
        self.assertEqual(text, "OCR text from image.")
        # Page count might be 0 if pdfplumber.open fails before counting pages
        self.assertEqual(page_count, 0) 
        mock_ocr_pdf.assert_called_once_with("bad.pdf")

    @patch('agents.document_extraction_agent.document_extraction_agent.pytesseract.image_to_string', side_effect=["OCR content from image 1.", "OCR content from image 2."])
    @patch('agents.document_extraction_agent.document_extraction_agent.convert_from_path')
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_ocr_pdf_success(self, mock_log_error, mock_convert_from_path, mock_image_to_string):
        """
        Test successful OCR extraction.
        """
        mock_image1 = MagicMock(spec=Image.Image)
        mock_image2 = MagicMock(spec=Image.Image)
        mock_convert_from_path.return_value = [mock_image1, mock_image2]

        result = self.agent._ocr_pdf("scanned.pdf")
        self.assertEqual(result, "OCR content from image 1.\n\nOCR content from image 2.\n\n")
        mock_convert_from_path.assert_called_once_with("scanned.pdf")
        self.assertEqual(mock_image_to_string.call_count, 2)
        mock_image_to_string.assert_any_call(mock_image1)
        mock_image_to_string.assert_any_call(mock_image2)
        mock_log_error.assert_not_called()

    @patch('agents.document_extraction_agent.document_extraction_agent.pytesseract.image_to_string', side_effect=Exception("Tesseract error"))
    @patch('agents.document_extraction_agent.document_extraction_agent.convert_from_path', return_value=[MagicMock(spec=Image.Image)])
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_ocr_pdf_failure(self, mock_log_error, mock_convert_from_path, mock_image_to_string):
        """
        Test OCR extraction failure.
        """
        result = self.agent._ocr_pdf("scanned.pdf")
        self.assertEqual(result, "")
        mock_convert_from_path.assert_called_once_with("scanned.pdf")
        mock_image_to_string.assert_called_once()
        
        mock_log_error.assert_called_once()
        call_args = mock_log_error.call_args[0]
        self.assertEqual(call_args[0], "scanned.pdf")
        self.assertEqual(call_args[1], "OCR extraction")
        self.assertIn("Tesseract error", call_args[2])

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text', return_value=("PDF content", 1))
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text')
    @patch('builtins.open', new_callable=mock_open)
    def test_process_file_pdf_success(self, mock_open, mock_extract_csv, mock_extract_pdf):
        """
        Test successful processing of a PDF file.
        """
        self.agent.process_file("file.pdf", "ID001", "VendorX", "time1")
        mock_extract_pdf.assert_called_once_with("file.pdf")
        mock_extract_csv.assert_not_called()
        
        mock_open.assert_called_with(self.agent.preprocessed_output_path, "a")
        written_content = mock_open().write.call_args[0][0]
        record = json.loads(written_content)
        self.assertEqual(record["source_id"], "ID001")
        self.assertEqual(record["format"], "pdf")
        self.assertEqual(record["text"], "PDF content")
        self.assertEqual(record["metadata"]["page_count"], 1)
        self.assertFalse(os.path.exists(self.agent.dead_letter_queue_path))
        self.assertFalse(os.path.exists(self.agent.unsupported_queue_path))

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text', return_value="CSV content")
    @patch('builtins.open', new_callable=mock_open)
    def test_process_file_csv_success(self, mock_open, mock_extract_csv, mock_extract_pdf):
        """
        Test successful processing of a CSV file.
        """
        self.agent.process_file("file.csv", "ID002", "VendorY", "time2")
        mock_extract_csv.assert_called_once_with("file.csv")
        mock_extract_pdf.assert_not_called()

        mock_open.assert_called_with(self.agent.preprocessed_output_path, "a")
        written_content = mock_open().write.call_args[0][0]
        record = json.loads(written_content)
        self.assertEqual(record["source_id"], "ID002")
        self.assertEqual(record["format"], "csv")
        self.assertEqual(record["text"], "CSV content")
        self.assertNotIn("page_count", record["metadata"])
        self.assertFalse(os.path.exists(self.agent.dead_letter_queue_path))
        self.assertFalse(os.path.exists(self.agent.unsupported_queue_path))

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text')
    def test_process_file_unsupported_type(self, mock_extract_csv, mock_extract_pdf):
        """
        Test processing of an unsupported file type.
        """
        self.agent.process_file("file.txt", "ID003", "VendorZ", "time3")
        mock_extract_csv.assert_not_called()
        mock_extract_pdf.assert_not_called()
        self.assertFalse(os.path.exists(self.agent.preprocessed_output_path))
        self.assertFalse(os.path.exists(self.agent.dead_letter_queue_path))
        with open(self.agent.unsupported_queue_path, "r") as f:
            log_content = f.read()
        expected_log = f"source_id: ID003, file_path: file.txt\n"
        self.assertIn(expected_log, log_content)

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text', return_value=("", 0))
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text')
    def test_process_file_no_text_extracted(self, mock_extract_csv, mock_extract_pdf):
        """
        Test processing when no text is extracted.
        """
        self.agent.process_file("empty.pdf", "ID004", "VendorA", "time4")
        self.assertFalse(os.path.exists(self.agent.preprocessed_output_path))
        self.assertFalse(os.path.exists(self.agent.unsupported_queue_path))
        with open(self.agent.dead_letter_queue_path, "r") as f:
            log_content = f.read()
        expected_log = f"source_id: ID004, step: text extraction, error: No text extracted from file.\n"
        self.assertIn(expected_log, log_content)

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text', side_effect=Exception("General error"))
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text')
    def test_process_file_general_error(self, mock_extract_csv, mock_extract_pdf):
        """
        Test general error handling during file processing.
        """
        self.agent.process_file("error.pdf", "ID005", "VendorB", "time5")
        self.assertFalse(os.path.exists(self.agent.preprocessed_output_path))
        self.assertFalse(os.path.exists(self.agent.unsupported_queue_path))
        with open(self.agent.dead_letter_queue_path, "r") as f:
            log_content = f.read()
        expected_log = f"source_id: ID005, step: overall processing, error: General error\n"
        self.assertIn(expected_log, log_content)

    @patch('agents.document_extraction_agent.document_extraction_agent.Document')
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_extract_docx_text_paragraphs(self, mock_log_error, mock_document):
        """
        Test DOCX extraction for paragraphs.
        """
        mock_doc_instance = MagicMock()
        mock_doc_instance.paragraphs = [
            MagicMock(text="Paragraph 1 content."),
            MagicMock(text="  Paragraph 2 content.  "),
            MagicMock(text=""), # Empty paragraph
            MagicMock(text="Paragraph 3 content.")
        ]
        mock_doc_instance.tables = []
        mock_document.return_value = mock_doc_instance

        expected_text = "Paragraph 1 content.\n\nParagraph 2 content.\n\nParagraph 3 content."
        result = self.agent._extract_docx_text("dummy.docx")
        self.assertEqual(result, expected_text)
        mock_log_error.assert_not_called()

    @patch('agents.document_extraction_agent.document_extraction_agent.Document')
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_extract_docx_text_tables(self, mock_log_error, mock_document):
        """
        Test DOCX extraction for tables with headers.
        """
        mock_doc_instance = MagicMock()
        mock_doc_instance.paragraphs = []

        # Mock a table with headers
        mock_table = MagicMock()
        # Mock rows as a list of MagicMocks, so list.index() works naturally
        mock_table.rows = [
            MagicMock(cells=[MagicMock(text="Header1"), MagicMock(text="Header2")]),
            MagicMock(cells=[MagicMock(text="Value1"), MagicMock(text="Value2")])
        ]
        
        mock_doc_instance.tables = [mock_table]
        mock_document.return_value = mock_doc_instance

        # Expected text after normalization in _extract_docx_text
        expected_text = "Header1: ; Header2: \nHeader1: Value1; Header2: Value2"
        result = self.agent._extract_docx_text("dummy.docx")
        self.assertEqual(result, expected_text)
        mock_log_error.assert_not_called()

    @patch('agents.document_extraction_agent.document_extraction_agent.Document')
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_extract_docx_text_empty(self, mock_log_error, mock_document):
        """
        Test DOCX extraction for an empty document.
        """
        mock_doc_instance = MagicMock()
        mock_doc_instance.paragraphs = []
        mock_doc_instance.tables = []
        mock_document.return_value = mock_doc_instance

        result = self.agent._extract_docx_text("empty.docx")
        self.assertEqual(result, "")
        mock_log_error.assert_not_called()

    @patch('agents.document_extraction_agent.document_extraction_agent.Document', side_effect=Exception("DOCX read error"))
    @patch.object(DocumentExtractionAgent, '_log_error')
    def test_extract_docx_text_failure(self, mock_log_error, mock_document):
        """
        Test DOCX extraction failure.
        """
        result = self.agent._extract_docx_text("bad.docx")
        self.assertEqual(result, "")
        mock_log_error.assert_called_once()
        call_args = mock_log_error.call_args[0]
        self.assertEqual(call_args[0], "bad.docx")
        self.assertEqual(call_args[1], "DOCX extraction")
        self.assertIn("DOCX read error", call_args[2])

    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_pdf_text')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_csv_text')
    @patch('agents.document_extraction_agent.document_extraction_agent.DocumentExtractionAgent._extract_docx_text', return_value="DOCX content")
    @patch('builtins.open', new_callable=mock_open)
    def test_process_file_docx_success(self, mock_open, mock_extract_docx, mock_extract_csv, mock_extract_pdf):
        """
        Test successful processing of a DOCX file.
        """
        self.agent.process_file("file.docx", "ID006", "VendorZ", "time6")
        mock_extract_docx.assert_called_once_with("file.docx")
        mock_extract_csv.assert_not_called()
        mock_extract_pdf.assert_not_called()

        mock_open.assert_called_with(self.agent.preprocessed_output_path, "a")
        written_content = mock_open().write.call_args[0][0]
        record = json.loads(written_content)
        self.assertEqual(record["source_id"], "ID006")
        self.assertEqual(record["format"], "docx")
        self.assertEqual(record["text"], "DOCX content")
        self.assertNotIn("page_count", record["metadata"])
        self.assertFalse(os.path.exists(self.agent.dead_letter_queue_path))
        self.assertFalse(os.path.exists(self.agent.unsupported_queue_path))

    def test_ocr_pdf_internal_logic(self):
        """
        Test the internal logic of _ocr_pdf, including success and failure scenarios.
        """
        # Test success
        with patch('agents.document_extraction_agent.document_extraction_agent.convert_from_path') as mock_convert_from_path, \
             patch('agents.document_extraction_agent.document_extraction_agent.pytesseract.image_to_string') as mock_image_to_string, \
             patch.object(DocumentExtractionAgent, '_log_error') as mock_log_error:
            
            mock_image1 = MagicMock(spec=Image.Image)
            mock_image2 = MagicMock(spec=Image.Image)
            mock_convert_from_path.return_value = [mock_image1, mock_image2]
            mock_image_to_string.side_effect = ["OCR content from image 1.", "OCR content from image 2."]

            result = self.agent._ocr_pdf("scanned_success.pdf")
            self.assertEqual(result, "OCR content from image 1.\n\nOCR content from image 2.\n\n")
            mock_convert_from_path.assert_called_once_with("scanned_success.pdf")
            self.assertEqual(mock_image_to_string.call_count, 2)
            mock_image_to_string.assert_any_call(mock_image1)
            mock_image_to_string.assert_any_call(mock_image2)
            mock_log_error.assert_not_called()
        
        # Test failure when convert_from_path fails
        self.setUp() # Reset mocks and agent state
        with patch('agents.document_extraction_agent.document_extraction_agent.convert_from_path', side_effect=Exception("Poppler error")) as mock_convert_from_path, \
             patch('agents.document_extraction_agent.document_extraction_agent.pytesseract.image_to_string') as mock_image_to_string, \
             patch.object(DocumentExtractionAgent, '_log_error') as mock_log_error:
            
            result = self.agent._ocr_pdf("scanned_poppler_fail.pdf")
            self.assertEqual(result, "")
            mock_convert_from_path.assert_called_once_with("scanned_poppler_fail.pdf")
            mock_image_to_string.assert_not_called()
            
            mock_log_error.assert_called_once()
            call_args = mock_log_error.call_args[0]
            self.assertEqual(call_args[0], "scanned_poppler_fail.pdf")
            self.assertEqual(call_args[1], "OCR extraction")
            self.assertIn("Poppler error", call_args[2])

        # Test failure when pytesseract.image_to_string fails
        self.setUp() # Reset mocks and agent state
        with patch('agents.document_extraction_agent.document_extraction_agent.convert_from_path', return_value=[MagicMock(spec=Image.Image)]) as mock_convert_from_path, \
             patch('agents.document_extraction_agent.document_extraction_agent.pytesseract.image_to_string', side_effect=Exception("Tesseract error")) as mock_image_to_string, \
             patch.object(DocumentExtractionAgent, '_log_error') as mock_log_error:
            
            result = self.agent._ocr_pdf("scanned_tesseract_fail.pdf")
            self.assertEqual(result, "")
            mock_convert_from_path.assert_called_once_with("scanned_tesseract_fail.pdf")
            mock_image_to_string.assert_called_once()
            
            mock_log_error.assert_called_once()
            call_args = mock_log_error.call_args[0]
            self.assertEqual(call_args[0], "scanned_tesseract_fail.pdf")
            self.assertEqual(call_args[1], "OCR extraction")
            self.assertIn("Tesseract error", call_args[2])

if __name__ == '__main__':
    unittest.main()
