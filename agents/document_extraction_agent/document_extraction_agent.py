import json
import logging
import re
import os
import pandas as pd
import pdfplumber
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Define a threshold for PDF text extraction to trigger OCR fallback
PDF_TEXT_THRESHOLD = 50  # characters

class DocumentExtractionAgent:
    """
    Agent 2: Document Extraction Agent
    Takes raw files (PDF or CSV), extracts and normalizes their text content,
    and emits a structured JSON record.
    """

    def __init__(self, preprocessed_output_path="preprocessed.jsonl", dead_letter_queue_path="dead_letter.log", unsupported_queue_path="unsupported_files.log"):
        self.preprocessed_output_path = preprocessed_output_path
        self.dead_letter_queue_path = dead_letter_queue_path
        self.unsupported_queue_path = unsupported_queue_path

    def _log_error(self, source_id: str, step: str, error: str):
        """Logs an error and sends the message to a dead-letter queue."""
        logging.error(f"Error processing {source_id} during {step}: {error}")
        with open(self.dead_letter_queue_path, "a") as f:
            f.write(f"source_id: {source_id}, step: {step}, error: {error}\n")

    def _log_unsupported(self, source_id: str, file_path: str):
        """Logs an unsupported file and sends the message to an unsupported queue."""
        logging.warning(f"Unsupported file type for {source_id}: {file_path}")
        with open(self.unsupported_queue_path, "a") as f:
            f.write(f"source_id: {source_id}, file_path: {file_path}\n")

    def _normalize_text(self, text: str) -> str:
        """
        Collapses all whitespace, fixes common encoding issues, and removes boilerplate.
        """
        if not text:
            return ""
        # Fix common encoding issues
        text = text.replace("â€™", "'").replace("â€œ", "\"").replace("â€ ", "\"").replace("â€", "\"")
        
        # Remove known boilerplate
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Confidential Document', '', text, flags=re.IGNORECASE)
        
        # Collapse multiple spaces or linebreaks to a single space
        text = re.sub(r'\s+', ' ', text)
        
        # Finally, strip leading/trailing whitespace
        text = text.strip()
        return text

    def _extract_pdf_text(self, file_path: str) -> tuple[str, int]:
        """
        Extracts text from PDF, with OCR fallback if text extraction is insufficient.
        Returns extracted text and page count.
        """
        extracted_text = ""
        page_count = 0
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n\n"

            if len(extracted_text.strip()) < PDF_TEXT_THRESHOLD:
                logging.info(f"Text-based PDF extraction too short or empty for {file_path}, falling back to OCR.")
                extracted_text = self._ocr_pdf(file_path)

        except Exception as e:
            logging.warning(f"Error during text-based PDF extraction for {file_path}: {e}. Falling back to OCR.")
            extracted_text = self._ocr_pdf(file_path)
        return extracted_text, page_count

    def _ocr_pdf(self, file_path: str) -> str:
        """
        Performs OCR on each page of a PDF.
        """
        ocr_text = ""
        try:
            images = convert_from_path(file_path)
            for i, image in enumerate(images):
                # Optional: Deskew/denoise image (requires libraries like OpenCV or scikit-image)
                # For simplicity, this example directly uses pytesseract
                page_ocr_text = pytesseract.image_to_string(image)
                ocr_text += page_ocr_text + "\n\n"
            logging.info(f"OCR completed for {file_path}.")
        except Exception as e:
            self._log_error(os.path.basename(file_path), "OCR extraction", str(e))
            return ""
        return ocr_text

    def _extract_csv_text(self, file_path: str) -> str:
        """
        Reads CSV file with pandas and converts each row into a text blob.
        """
        csv_text_blob = []
        try:
            df = pd.read_csv(file_path, dtype=str)  # Treat all columns as strings
            for index, row in df.iterrows():
                lines = []
                for col_name, value in row.items():
                    if pd.notna(value) and str(value).strip():  # Check for non-empty cells
                        lines.append(f"{col_name}: {value}")
                if lines:
                    csv_text_blob.append("\n".join(lines))
            return "\n\n".join(csv_text_blob)  # Concatenate multiple rows with separators
        except Exception as e:
            self._log_error(os.path.basename(file_path), "CSV extraction", str(e))
            return ""

    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extracts text from a DOCX file, including paragraphs and flattened tables.
        """
        full_text = []
        try:
            document = Document(file_path)
            for paragraph in document.paragraphs:
                normalized_paragraph_text = self._normalize_text(paragraph.text)
                if normalized_paragraph_text:
                    full_text.append(normalized_paragraph_text)
            
            for table in document.tables:
                table_text = []
                # Reason: Iterate through rows and cells to extract text and flatten tables.
                # The table flattening logic attempts to associate cell values with headers if available.
                for r_idx, row in enumerate(table.rows):
                    row_cells = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = self._normalize_text(cell.text)
                        if r_idx == 0: # Header row
                            if cell_text: # Only add if header is not empty
                                row_cells.append(f"{cell_text}: ")
                        else:
                            try:
                                header_cell_text = self._normalize_text(table.rows[0].cells[c_idx].text)
                                if header_cell_text:
                                    row_cells.append(f"{header_cell_text}: {cell_text}")
                                else:
                                    row_cells.append(cell_text)
                            except IndexError:
                                row_cells.append(cell_text)
                    if row_cells:
                        table_text.append("; ".join(row_cells))
                if table_text:
                    full_text.append("\n".join(table_text))
            
            return "\n\n".join(full_text)
        except Exception as e:
            self._log_error(os.path.basename(file_path), "DOCX extraction", str(e))
            return ""

    def process_file(self, file_path: str, source_id: str, vendor: str, timestamp: str):
        """
        Main function to process a single file.
        """
        logging.info(f"Starting processing for source_id: {source_id}, file: {file_path}")
        extracted_text = ""
        file_format = None
        page_count = None
        
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension == ".pdf":
                file_format = "pdf"
                extracted_text, page_count = self._extract_pdf_text(file_path)
            elif file_extension == ".csv":
                file_format = "csv"
                extracted_text = self._extract_csv_text(file_path)
            elif file_extension == ".docx":
                file_format = "docx"
                extracted_text = self._extract_docx_text(file_path)
            else:
                self._log_unsupported(source_id, file_path)
                return # Stop processing for unsupported types

            if not extracted_text:
                self._log_error(source_id, "text extraction", "No text extracted from file.")
                return

            normalized_text = self._normalize_text(extracted_text)

            output_record = {
                "source_id": source_id,
                "vendor": vendor,
                "format": file_format,
                "text": normalized_text,
                "metadata": {
                    "original_file": file_path,
                    "timestamp": timestamp,
                }
            }
            if page_count is not None:
                output_record["metadata"]["page_count"] = page_count

            # Append the JSON object as one line to preprocessed.jsonl
            with open(self.preprocessed_output_path, "a") as f:
                f.write(json.dumps(output_record) + "\n")
            
            logging.info(f"Successfully processed and emitted record for source_id: {source_id}")
            # Simulate publishing a "preprocessing_complete" event
            # In a real system, this would be a message queue publish
            logging.info(f"Event: preprocessing_complete for source_id: {source_id}")

        except Exception as e:
            self._log_error(source_id, "overall processing", str(e))

if __name__ == "__main__":
    # Example Usage (for testing purposes)
    agent = DocumentExtractionAgent()

    # Create dummy files for testing
    os.makedirs("temp_files", exist_ok=True)
    
    # Dummy PDF (requires a real PDF for full testing)
    # For actual testing, you'd need a PDF file.
    # Example: create a simple text file and convert to PDF for basic testing
    # Or use a pre-existing dummy.pdf
    with open("temp_files/dummy.txt", "w") as f:
        f.write("This is a test PDF document.\n\nIt has multiple lines.\n\nPage 2.")
    # You would typically use a tool like reportlab or fpdf to create a PDF from this text
    # For now, assume 'temp_files/dummy.pdf' exists or is created externally.
    # If you want to test OCR, ensure dummy.pdf is an image-based PDF.

    # Dummy CSV
    csv_content = """col1,col2,col3
value1,value2,value3
another_value,yet_another,final_value
"""
    with open("temp_files/dummy.csv", "w") as f:
        f.write(csv_content)

    # Dummy unsupported file
    with open("temp_files/dummy.txt", "w") as f:
        f.write("This is an unsupported text file.")

    # Process dummy files
    # Note: For PDF, you'll need to ensure pdfplumber and pytesseract can find the PDF and Tesseract executable.
    # This example assumes a 'dummy.pdf' exists for testing.
    # If you don't have a dummy.pdf, this part will likely fail unless you mock pdfplumber/pdf2image.
    # agent.process_file("temp_files/dummy.pdf", "INV_PDF_001", "VendorA", "2023-01-01T10:00:00Z")
    agent.process_file("temp_files/dummy.csv", "INV_CSV_001", "VendorB", "2023-01-01T10:05:00Z")
    agent.process_file("temp_files/dummy.txt", "INV_TXT_001", "VendorC", "2023-01-01T10:10:00Z")

    # Clean up dummy files
    # os.remove("temp_files/dummy.pdf") # Uncomment if you create a dummy.pdf
    os.remove("temp_files/dummy.csv")
    os.remove("temp_files/dummy.txt")
    os.rmdir("temp_files")

    print("\n--- preprocessed.jsonl content ---")
    if os.path.exists("preprocessed.jsonl"):
        with open("preprocessed.jsonl", "r") as f:
            print(f.read())
        os.remove("preprocessed.jsonl") # Clean up after example run
    
    print("\n--- dead_letter.log content ---")
    if os.path.exists("dead_letter.log"):
        with open("dead_letter.log", "r") as f:
            print(f.read())
        os.remove("dead_letter.log") # Clean up after example run

    print("\n--- unsupported_files.log content ---")
    if os.path.exists("unsupported_files.log"):
        with open("unsupported_files.log", "r") as f:
            print(f.read())
        os.remove("unsupported_files.log") # Clean up after example run
