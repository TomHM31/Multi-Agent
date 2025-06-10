import os
import json
import psycopg2
import mailparser
import pdfplumber
from docx import Document
import tiktoken
from sentence_transformers import SentenceTransformer
import boto3
from requests_aws4auth import AWS4Auth
import requests
import numpy as np

# Configuration
POSTGRES_CONN_STRING = "dbname=mydb user=postgres password=secret host=localhost"
EMAIL_DIR = "path/to/emails"
PDF_DIR = "path/to/pdfs"
DOCX_DIR = "path/to/docx"
OPENSEARCH_ENDPOINT = "your-opensearch-endpoint"
OPENSEARCH_INDEX = "company_docs"
AWS_REGION = "your-aws-region"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
MAX_TOKENS = 2000
OLLAMA_HOST = "http://localhost:11434"
OLLAMA_MODEL = "llama2"

class RAGSystem:
    def __init__(self):
        self.embed_model = SentenceTransformer(EMBEDDING_MODEL)
        self.enc = tiktoken.get_encoding("cl100k_base")
        self.session = boto3.Session()
        self.credentials = self.session.get_credentials()
        self.awsauth = AWS4Auth(
            self.credentials.access_key,
            self.credentials.secret_key,
            AWS_REGION,
            "es",
            session_token=self.credentials.token
        )
    
    def extract_sql_data(self):
        """Extract data from SQL tables and convert to text records"""
        records = []
        try:
            conn = psycopg2.connect(POSTGRES_CONN_STRING)
            cursor = conn.cursor()
            
            # Example: Extract workflow steps
            cursor.execute("SELECT * FROM workflow_steps")
            for row in cursor.fetchall():
                record = {
                    "source": "workflow_steps",
                    "id": f"workflow_{row[0]}",
                    "text": f"Workflow {row[1]} - Step {row[2]}: {row[3]}; Role: {row[4]}",
                    "metadata": {
                        "workflow_name": row[1],
                        "step_order": row[2],
                        "description": row[3],
                        "role": row[4]
                    }
                }
                records.append(record)
            
            # Add extraction for other tables as needed
            cursor.close()
            conn.close()
        except Exception as e:
            print(f"Error extracting SQL data: {e}")
        return records
    
    def extract_email_data(self):
        """Extract data from email files"""
        records = []
        for filename in os.listdir(EMAIL_DIR):
            if filename.endswith(".eml") or filename.endswith(".msg"):
                try:
                    mail = mailparser.parse_from_file(os.path.join(EMAIL_DIR, filename))
                    record = {
                        "source": "email",
                        "id": f"email_{mail.date.strftime('%Y-%m-%d')}_{filename}",
                        "text": f"From: {mail.from_}\nDate: {mail.date}\nSubject: {mail.subject}\n\n{mail.body}",
                        "metadata": {
                            "subject": mail.subject,
                            "sender": mail.from_,
                            "date": mail.date.strftime("%Y-%m-%d")
                        }
                    }
                    records.append(record)
                except Exception as e:
                    print(f"Error processing email {filename}: {e}")
        return records
    
    def extract_pdf_data(self):
        """Extract data from PDF documents"""
        records = []
        for filename in os.listdir(PDF_DIR):
            if filename.endswith(".pdf"):
                try:
                    with pdfplumber.open(os.path.join(PDF_DIR, filename)) as pdf:
                        for i, page in enumerate(pdf.pages):
                            text = page.extract_text()
                            if text:
                                record = {
                                    "source": "pdf",
                                    "id": f"{os.path.splitext(filename)[0]}_page_{i+1}",
                                    "text": text,
                                    "metadata": {
                                        "filename": filename,
                                        "page": i+1
                                    }
                                }
                                records.append(record)
                except Exception as e:
                    print(f"Error processing PDF {filename}: {e}")
        return records
    
    def extract_docx_data(self):
        """Extract data from DOCX files"""
        records = []
        for filename in os.listdir(DOCX_DIR):
            if filename.endswith(".docx"):
                try:
                    doc = Document(os.path.join(DOCX_DIR, filename))
                    for i, para in enumerate(doc.paragraphs):
                        if para.text.strip():
                            record = {
                                "source": "docx",
                                "id": f"{os.path.splitext(filename)[0]}_para_{i+1}",
                                "text": para.text,
                                "metadata": {
                                    "filename": filename,
                                    "para": i+1
                                }
                            }
                            records.append(record)
                    
                    # Extract tables
                    for t, table in enumerate(doc.tables):
                        table_text = ""
                        for row in table.rows:
                            row_text = "\t".join(cell.text for cell in row.cells)
                            table_text += row_text + "\n"
                        
                        if table_text:
                            record = {
                                "source": "docx_table",
                                "id": f"{os.path.splitext(filename)[0]}_table_{t+1}",
                                "text": table_text,
                                "metadata": {
                                    "filename": filename,
                                    "table": t+1
                                }
                            }
                            records.append(record)
                except Exception as e:
                    print(f"Error processing DOCX {filename}: {e}")
        return records
    
    def chunk_text(self, records):
        """Chunk text records into smaller pieces"""
        chunks = []
        for record in records:
            tokens = self.enc.encode(record["text"])
            chunk_index = 1
            while tokens:
                piece = tokens[:MAX_TOKENS]
                chunk_text = self.enc.decode(piece)
                
                chunk = {
                    "chunk_id": f"{record['id']}_chunk_{chunk_index}",
                    "text": chunk_text,
                    "metadata": {
                        **record["metadata"],
                        "source": record["source"],
                        "original_id": record["id"],
                        "chunk_index": chunk_index
                    }
                }
                chunks.append(chunk)
                
                tokens = tokens[MAX_TOKENS:]
                chunk_index += 1
        return chunks
    
    def create_opensearch_index(self):
        """Create the OpenSearch index with k-NN mapping"""
        url = f"https://{OPENSEARCH_ENDPOINT}/{OPENSEARCH_INDEX}"
        mapping = {
            "settings": {
                "index.knn": True,
                "index.knn.space_type": "l2"
            },
            "mappings": {
                "properties": {
                    "chunk_id": {"type": "keyword"},
                    "text": {"type": "text"},
                    "embedding": {"type": "knn_vector", "dimension": 384},
                    "source": {"type": "keyword"},
                    "original_id": {"type": "keyword"},
                    "chunk_index": {"type": "integer"}
                }
            }
        }
        
        response = requests.put(url, auth=self.awsauth, json=mapping, headers={"Content-Type": "application/json"})
        if response.status_code not in [200, 201]:
            print(f"Failed to create index: {response.text}")
        else:
            print(f"Index {OPENSEARCH_INDEX} created successfully")
    
    def index_chunks(self, chunks):
        """Compute embeddings and index chunks into OpenSearch"""
        url = f"https://{OPENSEARCH_ENDPOINT}/{OPENSEARCH_INDEX}/_doc"
        
        for chunk in chunks:
            embedding = self.embed_model.encode([chunk["text"]], convert_to_numpy=True)[0].tolist()
            document = {
                "chunk_id": chunk["chunk_id"],
                "text": chunk["text"],
                "embedding": embedding,
                "source": chunk["metadata"]["source"],
                "original_id": chunk["metadata"]["original_id"],
                "chunk_index": chunk["metadata"]["chunk_index"]
            }
            
            response = requests.put(
                f"{url}/{chunk['chunk_id']}",
                auth=self.awsauth,
                json=document,
                headers={"Content-Type": "application/json"}
            )
            
            if response.status_code not in [200, 201]:
                print(f"Failed to index chunk {chunk['chunk_id']}: {response.text}")
    
    def call_llm(self, prompt, max_tokens=512, temperature=0.0, stop=None):
        """Call the LLM via Ollama"""
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        if stop:
            payload["stop"] = stop
            
        try:
            response = requests.post(
                f"{OLLAMA_HOST}/v1/completions",
                json=payload,
                timeout=60
            )
            response.raise_for_status()
            return response.json()["choices"][0]["text"]
        except Exception as e:
            print(f"Error calling LLM: {e}")
            return None
    
    def search_opensearch(self, query_vector, k=4):
        """Search OpenSearch for top-k relevant chunks"""
        url = f"https://{OPENSEARCH_ENDPOINT}/{OPENSEARCH_INDEX}/_search"
        query = {
            "size": k,
            "query": {
                "knn": {
                    "embedding": {
                        "vector": query_vector,
                        "k": k
                    }
                }
            }
        }
        
        response = requests.post(
            url,
            auth=self.awsauth,
            json=query,
            headers={"Content-Type": "application/json"}
        )
        
        if response.status_code == 200:
            return response.json()["hits"]["hits"]
        else:
            print(f"Search failed: {response.text}")
            return []
    
    def query_rag(self, question):
        """Retrieve relevant information and generate an answer"""
        # Embed the question
        q_vec = self.embed_model.encode([question], convert_to_numpy=True)[0].tolist()
        
        # Retrieve relevant chunks
        results = self.search_opensearch(q_vec)
        if not results:
            return "No relevant information found."
        
        # Build the prompt
        prompt = "Below are relevant excerpts from our company knowledge base:\n\n"
        for i, hit in enumerate(results):
            source = hit["_source"]
            prompt += f"--- Chunk {i+1} (id: {source['chunk_id']}) ---\n"
            prompt += f"{source['text']}\n\n"
        
        prompt += f"Question: {question}\n"
        prompt += "Please answer using only the passages above."
        
        # Generate the answer
        return self.call_llm(prompt, max_tokens=512, temperature=0.0, stop=["\n\n"])
    
    def run_full_pipeline(self):
        """Run the full RAG pipeline"""
        print("Extracting data from sources...")
        sql_data = self.extract_sql_data()
        email_data = self.extract_email_data()
        pdf_data = self.extract_pdf_data()
        docx_data = self.extract_docx_data()
        
        all_data = sql_data + email_data + pdf_data + docx_data
        print(f"Extracted {len(all_data)} records from all sources")
        
        print("Chunking records...")
        chunks = self.chunk_text(all_data)
        print(f"Created {len(chunks)} chunks")
        
        print("Creating OpenSearch index...")
        self.create_opensearch_index()
        
        print("Indexing chunks...")
        self.index_chunks(chunks)
        print("Indexing completed")
        
        # Example query
        print("\nTesting RAG system...")
        question = "What is our standard invoice-approval process?"
        answer = self.query_rag(question)
        print(f"Question: {question}")
        print(f"Answer: {answer}")

if __name__ == "__main__":
    rag_system = RAGSystem()
    rag_system.run_full_pipeline()
